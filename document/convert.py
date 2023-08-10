import importlib
import sys
import xlrd
from docx import Document

def main():
    importlib.reload(sys)
    # sys.setdefaultencoding('utf-8') python3.10 sys has no method setdefaultencoding

    fileName = 'fileName'
    excelPath = './'+fileName+'.xlsx'
    docPath = './'+ fileName +'.docx'
    sheetName = 'Sheet1'
    data = xlrd.open_workbook(excelPath)
    table = data.sheet_by_name(sheetName)

    # row amout 
    rowAmount = table.nrows
    # col amount 
    colAmount = table.ncols
    list = ("单选题","多选题","判断题","填空题")
    print(rowAmount,colAmount)
    
    singleChoice = "单选题"
    multipleChoice = "多选题"
    yesOrNo = "判断题"
    gapFilling = "填空题"


    # create document
    document = Document()
    
    seqNo = 1
    for rowIndex in range(rowAmount):
        if table.cell_value(rowIndex, 1) == singleChoice:
            question =  table.cell_value(rowIndex, 2)
            a = table.cell_value(rowIndex, 3)
            b = table.cell_value(rowIndex, 4)
            c = table.cell_value(rowIndex, 5)
            d = table.cell_value(rowIndex, 6)
            insertSingleChoice(document=document,seqNo=seqNo,question=question,a=a,b=b,c=c,d=d)
            seqNo += 1  

    seqNo = 1
    for rowIndex in range(rowAmount):
        if table.cell_value(rowIndex, 1) == multipleChoice:
            question =  table.cell_value(rowIndex, 2)
            a = table.cell_value(rowIndex, 3)
            b = table.cell_value(rowIndex, 4)
            c = table.cell_value(rowIndex, 5)
            d = table.cell_value(rowIndex, 6)
            insertMultiChoice(document=document,seqNo=seqNo,question=question,a=a,b=b,c=c,d=d)
            seqNo += 1    
    seqNo = 1
    for rowIndex in range(rowAmount):
        if table.cell_value(rowIndex, 1) == yesOrNo:
            question =  table.cell_value(rowIndex, 2)
            insertYesOrNo(document=document,seqNo=seqNo,question=question)
            seqNo += 1    
    seqNo = 1
    for rowIndex in range(rowAmount):
        if table.cell_value(rowIndex, 1) == gapFilling:
            question =  table.cell_value(rowIndex, 2)
            insertGapFilling(document=document,seqNo=seqNo,question=question)
            seqNo += 1    

                  
    # 保存文档
    document.save(docPath)

def insertSingleChoice(document,seqNo,question,a,b,c,d):
    if seqNo == 1:
        document.add_heading('一、单项选择题',level=1)
    
    # add options
    document.add_paragraph(str(seqNo)+'.'+question)
    document.add_paragraph('A.'+str(a))
    document.add_paragraph('B.'+str(b))
    document.add_paragraph('C.'+str(c))
    document.add_paragraph('D.'+str(d))

def insertMultiChoice(document,seqNo,question,a,b,c,d):
    if seqNo == 1:
        document.add_heading('二、多项选择题',level=1)
    # add options
    document.add_paragraph(str(seqNo)+'.'+question)
    document.add_paragraph('A.'+str(a))
    document.add_paragraph('B.'+str(b))
    document.add_paragraph('C.'+str(c))
    document.add_paragraph('D.'+str(d))

def insertYesOrNo(document,seqNo,question):
    if seqNo == 1:
        document.add_heading('三、判断题',level=1)
    document.add_paragraph(str(seqNo)+'.'+question)

def  insertGapFilling(document,seqNo,question):
    if seqNo == 1:
        document.add_heading('四、填空题',level=1)
    document.add_paragraph(str(seqNo)+'.'+question)


if __name__ == '__main__':
    main()